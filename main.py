import sys, logging, time, re, io
from datetime import datetime

from docx import Document
from htmldocx import HtmlToDocx

import class_read_csv
csv_import_class = class_read_csv.Read_CSV_Data()

# Random
today_date = datetime.today().strftime("%d-%m-%Y").replace("-", "")

# Setup Logging
log_format = (
    "%(asctime)s::%(levelname)s::%(name)s::" "%(filename)s::%(lineno)d::%(message)s"
)
logging.basicConfig(
    filename="logs_" + today_date + ".log", level="INFO", format=log_format
)

inclusion_list = {"Versioning for Admin View": ["GAPS-17"],
                  "Content Expiry Reporting": ["GAPS-20", "GAPS-83", "GAPS-89", "GAPS-88", "GAPS-93", "GAPS-25",
                                               "GAPS-30", "GAPS-87", "GAPS-24", "GAPS-26", "GAPS-29", "GAPS-23",
                                                "GAPS-27", "GAPS-21", "GAPS-22", "GAPS-28", "GAPS-90", "GAPS-91",
                                                "GAPS-92"],
                  "Compliance Workflows": ["GAPS-36", "GAPS-37", "GAPS-38", "GAPS-39", "GAPS-40"],
                  "UI/UX Improvements to Django interface": ["GAPS-48", "GAPS-49", "GAPS-50", "GAPS-97", "GAPS-51",
                                                            "GAPS-52", "GAPS-53", "GAPS-54", "GAPS-9", "GAPS-47"],
                  "CMS 4.0 Regressions": ["GAPS-56", "GAPS-57",
                  "GAPS-59", "GAPS-60", "GAPS-61", "GAPS-62", "GAPS-63", "GAPS-64", "GAPS-66", "GAPS-67", "GAPS-68",
                  "GAPS-69", "GAPS-70", "GAPS-71", "GAPS-72", "GAPS-73", "GAPS-74", "GAPS-79", "GAPS-80", "GAPS-81"],
                  "Page Admin Enhancements": ["GAPS-31", "GAPS-33", "GAPS-34", "GAPS-45", "GAPS-46"]

                  }
# Rules taken from: https://jira.atlassian.com/secure/WikiRendererHelpAction.jspa?section=all
bold_rule = re.compile('(\*.*?\*)')
heading_rule = re.compile('(h[1-6].*)')


def create_document(csv_data):

    try:
        text_file = open("items.txt", "w")
        html_file = open("demofile.html", "w")
        html_file.write("<!DOCTYPE html>\n")
        html_file.write("<html>\n")
        html_file.write("<body>\n")
        row_data = {}
        document = Document()
        document.add_heading('Gaps Items', 0)

        for key, val in inclusion_list.items():
            html_file.write(f"<h1> {key} </h1>\n")
            document.add_paragraph(key, style="Heading 1")
            for a in val:

                for i in csv_data:
                    if i["Issue key"] == a:
                        row_data["Issue key"] = check_string(i["Issue key"])
                        row_data["Summary"] = check_string(i["Summary"])
                        row_data["Description"] = check_string(i["Description"])

                create_text_file(row_data, text_file, document, html_file)
        html_file.write("</body>\n")
        html_file.write("</html>\n")
        new_parser = HtmlToDocx()
        new_parser.parse_html_file("demofile.html", "gaps_items_html_converted")

    except Exception as e:
        logging.error(str(e))


def create_text_file(row_data, text_obj, doc_obj, html_file):

    try:
        # doc_obj.add_paragraph(row_data["Issue key"] + ": " + row_data["Summary"], style="Heading 3")
        # for line in io.StringIO(row_data["Description"]):
        #     doc_obj.add_paragraph(replace_bold(replace_heading(line)), style="Body Text")
        # text_obj.write(f"{row_data['Issue key']} : {row_data['Summary']} \n")
        # text_obj.write(row_data["Description"] + "\n")
        # text_obj.write("\n")
        # doc_obj.save('Gaps_items.docx')

        html_file.write(f"<h3> {row_data['Issue key']} : {row_data['Summary']} </h3>\n")
        for line in io.StringIO(row_data["Description"]):
            html_file.write(f"<p> {replace_ordered_list(replace_unordered_list(replace_image(replace_bold(replace_heading(line)))))} </p>\n")
    except Exception as e:
        logging.error(str(e))


def check_string(value):
    if len(value) == 0:
        return "N/a"
    else:
        return value


def replace_image(source):
    if source[0] == "!":
        return "Image"
    else:
        return source


def proccess_bullets(line, bullet_state, list_tag_type="ul"):
    line_replacement = ""

    def _open_tag(tag):
        return f"<{tag}>\n"

    def _enclosed_tag(tag, string):
        return f"<{tag}>{string}</{tag}>\n"

    def _close_tag(tag):
        return f"</{tag}>\n"

    # Current item is a bullet point
    if bullet_state['current_level']:
        # Open a new level
        if not bullet_state['last_level'] or bullet_state['current_level'] > bullet_state['last_level']:
            line_replacement += _open_tag(list_tag_type)
        # Close existing level i.e the previous indentation has ended
        elif bullet_state['last_level'] > bullet_state['current_level']:
            line_replacement += _close_tag(list_tag_type)

        cleaned_line = line[len(bullet_state['prefix']):]
        line_replacement += _enclosed_tag("li", cleaned_line)
    # Current item is not a bullet point, may need to close off existing levels
    else:
        if bullet_state['last_level']:
            # close levels
            for level in range(0, bullet_state['last_level']):
                line_replacement += _close_tag(list_tag_type)

        line_replacement += line

    bullet_state['last_level'] = bullet_state['current_level']
    return line_replacement


def replace_unordered_list(source):
    # Sentences starting with: "#", "##", "###" and so on for an ordered list 1, 1.1, 2, 2.2, 2.3, 3
    lines = io.StringIO(source)
    bullet_state = {
        'current_level': False,
        'last_level': False,
        'prefix': "",
    }

    def _calculate_bullet_level(line):
        if line.startswith("* "):
            return 1, "* "
        elif line.startswith("** "):
            return 2, "** "
        elif line.startswith("*** "):
            return 3, "*** "
        elif line.startswith("**** "):
            return 4, "**** "
        elif line.startswith("***** "):
            return 5, "***** "
        elif line.startswith("****** "):
            return 6, "****** "
        return False, ""

    output = ""
    for line in lines:
        bullet_state['current_level'], bullet_state['prefix'] = _calculate_bullet_level(line)
        output += proccess_bullets(line, bullet_state)

    return output


def replace_ordered_list(source):
    # Sentences starting with: "#", "##", "###" and so on for an ordered list 1, 1.1, 2, 2.2, 2.3, 3
    lines = io.StringIO(source)
    bullet_state = {
        'current_level': False,
        'last_level': False,
        'prefix': "",
    }

    def _calculate_bullet_level(line):
        if line.startswith("# "):
            return 1, "# "
        elif line.startswith("## "):
            return 2, "## "
        elif line.startswith("### "):
            return 3, "### "
        elif line.startswith("#### "):
            return 4, "#### "
        elif line.startswith("##### "):
            return 5, "##### "
        elif line.startswith("###### "):
            return 6, "###### "
        return False, ""

    output = ""
    for line in lines:
        bullet_state['current_level'], bullet_state['prefix'] = _calculate_bullet_level(line)
        output += proccess_bullets(line, bullet_state, list_tag_type="ol")

    return output


def replace_bold(source):
    # Characters wrapped in *strong*
    for match in bold_rule.findall(source):
        cleaned_string = match.replace("*", "")
        # TODO: Write change to file
        source = source.replace(match, cleaned_string)
    return source


def replace_heading(source):
    # Words starting with: "h1.", "h2.", "h3.", "h4.", "h5." or "h6."
    for match in heading_rule.findall(source):
        heading_number = match[1]
        cleaned_string = match.replace(f"h{heading_number}. ", "")
        source = source.replace(match, cleaned_string)
        # TODO: Write change to file
    return source

def replace_numbered(source):
    # Words starting with: "#"
    for match in heading_rule.findall(source):
        heading_number = match[1]
        cleaned_string = match.replace(f"h{heading_number}. ", "")
        source = source.replace(match, cleaned_string)
        # TODO: Write change to file
    return source


if __name__ == "__main__":
    csv_data = csv_import_class.read_csv("gaps.csv")
    create_document(csv_data)